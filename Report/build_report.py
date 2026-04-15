"""Build the full FYP report DOCX for the Smart Attendance System project.

Uses python-docx to produce a ~80-100 page report matching the GIK Institute
senior design project template.
"""
import os
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION, WD_ORIENTATION

HERE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(HERE, "assets")
OUT = os.path.join(HERE, "FYP_Report.docx")
LOGO = os.path.join(ASSETS, "giki_logo.jpeg")


# =============================================================================
# Helper functions
# =============================================================================

def _set_spacing(paragraph, before=0, after=6, line=1.15):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def set_cell_shading(cell, hex_color):
    """Set background color on a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_toc_field(paragraph):
    """Insert a Word TOC field that auto-populates when opened in Word."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    fldChar1.set(qn("w:dirty"), "true")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "Right-click and select Update Field to populate Table of Contents."
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    run._r.append(fldChar4)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
        if level == 0:
            run.font.size = Pt(22)
        elif level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
        run.bold = True
    return h


def add_para(doc, text, size=11, align=None, before=0, after=6, line=1.15, bold=False,
             italic=False, color=None, font="Times New Roman"):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color is not None:
        run.font.color.rgb = color
    _set_spacing(p, before=before, after=after, line=line)
    return p


def add_bullets(doc, items, font="Times New Roman"):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.runs[0] if p.runs else p.add_run()
        if p.runs:
            run.text = ""
        r = p.add_run(item)
        r.font.size = Pt(11)
        r.font.name = font
        _set_spacing(p, before=0, after=3, line=1.15)


def add_numbered(doc, items, font="Times New Roman"):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        r.font.size = Pt(11)
        r.font.name = font
        _set_spacing(p, before=0, after=3, line=1.15)


def add_figure(doc, filename, caption, width_inches=6.0):
    if not os.path.exists(os.path.join(ASSETS, filename)):
        print(f"WARNING: missing figure {filename}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(os.path.join(ASSETS, filename), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(10)
    r.font.name = "Times New Roman"
    _set_spacing(cap, before=3, after=12, line=1.15)


def add_table(doc, headers, rows, caption=None, col_widths=None):
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = "Times New Roman"
        _set_spacing(cap, before=6, after=3, line=1.15)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[i], "1F3A68")
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_i, row_data in enumerate(rows):
        row_cells = table.rows[r_i + 1].cells
        for c_i, val in enumerate(row_data):
            row_cells[c_i].text = ""
            p = row_cells[c_i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"
            row_cells[c_i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for col_idx, width in enumerate(col_widths):
            for cell in table.columns[col_idx].cells:
                cell.width = Inches(width)

    doc.add_paragraph()  # spacing after
    return table


def add_page_break(doc):
    doc.add_page_break()


def add_section_break(doc):
    doc.add_section(WD_SECTION.NEW_PAGE)


# =============================================================================
# Document setup
# =============================================================================

def setup_document():
    doc = Document()

    # Page margins: 1 inch on all sides
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    return doc


# =============================================================================
# Cover Page
# =============================================================================

def build_cover(doc):
    TITLE = "Smart Attendance System using Facial Image Recognition"

    # Blank line at top
    doc.add_paragraph()

    # Title
    add_para(doc, TITLE, size=24, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True,
             color=RGBColor(0x1F, 0x3A, 0x68), font="Arial", after=12, before=0)

    # Subtitle
    add_para(doc, "Senior Design Project (Part I & II)", size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, font="Arial", after=4)

    add_para(doc, "", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)

    # Submitted by
    add_para(doc, "Submitted by:", size=13, align=WD_ALIGN_PARAGRAPH.CENTER,
             bold=True, after=10)

    # Authors table
    members = [
        ("Hassaan Ahmed", "FCSE"),
        ("Osaid Khan Afridi", "FCSE"),
        ("Saad Khan", "FEE"),
        ("Najaf", "FEE"),
    ]
    table = doc.add_table(rows=len(members), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (name, dept) in enumerate(members):
        row = table.rows[i].cells
        row[0].text = ""
        p = row[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        row[1].text = ""
        p = row[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(dept)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        row[0].width = Inches(3.0)
        row[1].width = Inches(1.5)

    add_para(doc, "", size=10, after=20)

    # Advisor + Co-advisor
    add_para(doc, "Advisor: Dr. Zaiwar Ali", size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, after=4)
    add_para(doc, "Co-advisor: Dr. Arbab Abdur Rahim", size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, after=30)

    # Logo
    if os.path.exists(LOGO):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(LOGO, width=Inches(1.8))

    add_para(doc, "", size=8, after=10)
    add_para(doc, "Faculty of Electrical Engineering", size=13,
             align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, after=2)
    add_para(doc, "GIK Institute of Engineering Sciences and Technology", size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, after=2)
    add_para(doc, "Topi, District Swabi, Khyber Pakhtunkhwa, Pakistan",
             size=11, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, after=2)
    add_para(doc, "www.giki.edu.pk", size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, after=12)

    add_para(doc, "April 2026", size=13,
             align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, after=4)

    add_page_break(doc)


# =============================================================================
# Dedication
# =============================================================================

def build_dedication(doc):
    add_para(doc, "", size=14, after=300, before=200)
    add_para(doc, "Dedicated to our dear homeland, Pakistan, and to the",
             size=18, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, bold=True,
             color=RGBColor(0x1F, 0x3A, 0x68), font="Georgia")
    add_para(doc, "teachers, families, and friends who made this journey possible.",
             size=18, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, bold=True,
             color=RGBColor(0x1F, 0x3A, 0x68), font="Georgia", after=20)
    add_page_break(doc)


# =============================================================================
# Certificate of Ownership
# =============================================================================

def build_certificate(doc):
    add_heading_styled(doc, "Certificate of Ownership", level=1)
    add_para(doc, "", after=12)

    add_para(doc,
             "It is certified that the work contained in this final year project report "
             "titled:",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=12)

    add_para(doc, "\u201CSmart Attendance System using Facial Image Recognition\u201D",
             size=13, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, italic=True,
             color=RGBColor(0x1F, 0x3A, 0x68), after=12)

    add_para(doc,
             "has been carried out through the collective efforts of "
             "Hassaan Ahmed, Osaid Khan Afridi, Saad Khan, and Najaf under the "
             "supervision of Dr. Zaiwar Ali and the co-supervision of "
             "Dr. Arbab Abdur Rahim, and that to the best of our knowledge this work "
             "has not been submitted elsewhere for the award of any academic degree.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=20)

    # Advisor signature block
    add_para(doc, "_" * 40, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, after=2)
    add_para(doc, "Advisor: Dr. Zaiwar Ali", size=11, bold=True, after=2)
    add_para(doc, "Faculty of Electrical Engineering", size=11, after=2)
    add_para(doc, "GIK Institute of Engineering Sciences and Technology",
             size=11, italic=True, after=20)

    # Co-advisor signature block
    add_para(doc, "_" * 40, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, after=2)
    add_para(doc, "Co-advisor: Dr. Arbab Abdur Rahim", size=11, bold=True, after=2)
    add_para(doc, "Dean, Faculty of Electrical Engineering", size=11, after=2)
    add_para(doc, "GIK Institute of Engineering Sciences and Technology",
             size=11, italic=True, after=30)

    add_para(doc, "Date of submission: 15th April 2026",
             size=11, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, after=12)

    add_page_break(doc)


# =============================================================================
# Preface
# =============================================================================

def build_preface(doc):
    add_heading_styled(doc, "Preface", level=1)
    add_para(doc, "", after=4)

    paragraphs = [
        "Attendance in universities has historically been handled through roll-call or "
        "paper registers \u2014 a workflow that consumes classroom time, is prone to "
        "proxy attendance, and produces data that is hard to audit or aggregate at a "
        "departmental level. At the same time, smartphones have made it trivial to "
        "capture a high-resolution classroom photograph in seconds. This juxtaposition "
        "motivated the present project: can a simple classroom photograph replace the "
        "manual attendance register, while producing more accurate, more tamper-resistant "
        "records?",

        "The system described in this report answers that question. It is a production-style "
        "full-stack application consisting of a React frontend for teachers and "
        "administrators, a Node.js/TypeScript backend with role-based authentication and a "
        "Prisma-managed SQLite database, and a Python FastAPI AI service that performs face "
        "detection (MTCNN with a RetinaFace fallback) and face recognition (FaceNet 512-d "
        "embeddings followed by a custom 128-d projection head trained with triplet loss, "
        "backed by a support-vector-machine classifier). The recognition pipeline achieves "
        "86% recognition rate on unseen classroom group photographs across 14 test sessions "
        "and 243 faces drawn from two student batches of the GIK Institute.",

        "For readers who study full-stack or machine-learning systems, this report documents "
        "not just the final state of the system but also the design decisions, ablations, "
        "and lessons learned along the way \u2014 including why a cosine-distance ArcFace "
        "pipeline was abandoned in favour of a Euclidean-distance FaceNet pipeline, how a "
        "small domain-adapted projection head bridges the gap between controlled selfie "
        "enrollment and uncontrolled classroom photographs, and why quality filters and "
        "margin-based confidence scoring matter for correctness as much as raw accuracy "
        "numbers.",
    ]
    for p in paragraphs:
        add_para(doc, p, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=10, line=1.3)

    add_page_break(doc)


# =============================================================================
# Acknowledgments
# =============================================================================

def build_acknowledgments(doc):
    add_heading_styled(doc, "Acknowledgments", level=1)

    add_para(doc,
             "We are grateful to Almighty Allah for providing us with knowledge, courage, "
             "patience, and strength throughout this project, and for the continuous support "
             "of our families, without which this work would not have been possible.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=10, line=1.3)

    add_para(doc,
             "We extend our deepest gratitude to our advisor, Dr. Zaiwar Ali, whose "
             "guidance, availability, and technical insights consistently steered the project "
             "in the right direction. His willingness to engage with the details \u2014 from "
             "choice of face-recognition backbone to the design of our enrollment workflow "
             "\u2014 turned what could have been a purely academic exercise into an end-to-end "
             "system we could defend with data.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=10, line=1.3)

    add_para(doc,
             "We are similarly indebted to our co-advisor, Dr. Arbab Abdur Rahim, Dean of "
             "the Faculty of Electrical Engineering, for his support, administrative guidance, "
             "and the resources he made available at critical points during the project.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=10, line=1.3)

    add_para(doc,
             "We also thank the many students of the 2022 and 2023 batches who agreed to "
             "contribute their selfie photographs and classroom images to our dataset, and "
             "the teaching staff of the Faculty of Electrical Engineering for their "
             "feedback during early user-testing sessions. Finally, we thank GIK Institute of "
             "Engineering Sciences and Technology for providing us with the environment, "
             "laboratories, and academic freedom to pursue this work.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=16, line=1.3)

    add_para(doc, "Hassaan Ahmed, April 2026", size=11, italic=True, after=4)
    add_para(doc, "Osaid Khan Afridi, April 2026", size=11, italic=True, after=4)
    add_para(doc, "Saad Khan, April 2026", size=11, italic=True, after=4)
    add_para(doc, "Najaf, April 2026", size=11, italic=True, after=10)

    add_page_break(doc)


# =============================================================================
# TOC, LOF, LOT
# =============================================================================

def build_toc(doc):
    add_heading_styled(doc, "Table of Contents", level=1)
    p = doc.add_paragraph()
    add_toc_field(p)
    add_page_break(doc)


def build_lof(doc):
    add_heading_styled(doc, "List of Figures", level=1)
    figs = [
        ("Figure 3.1", "High-level system architecture"),
        ("Figure 3.2", "Simplified entity-relationship diagram"),
        ("Figure 3.3", "Attendance session data flow"),
        ("Figure 4.1", "Face recognition pipeline"),
        ("Figure 4.2", "Projection head training curves"),
        ("Figure 4.3", "Intra- vs inter-class distance distribution"),
        ("Figure 5.1", "Threshold ablation (recall vs false-positive rate)"),
        ("Figure 5.2", "Model comparison: ArcFace vs FaceNet vs FaceNet + projection"),
        ("Figure 5.3", "Recognition performance across test photos"),
        ("Figure 1.1", "Project Gantt chart"),
    ]
    for num, title in figs:
        p = doc.add_paragraph()
        run = p.add_run(f"{num}\t")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
        run2 = p.add_run(title)
        run2.font.size = Pt(11)
        run2.font.name = "Times New Roman"
        _set_spacing(p, before=0, after=3, line=1.2)
    add_page_break(doc)


def build_lot(doc):
    add_heading_styled(doc, "List of Tables", level=1)
    tabs = [
        ("Table 1.1", "Contribution of team members"),
        ("Table 2.1", "Comparison of surveyed face-recognition backbones"),
        ("Table 3.1", "Core database entities"),
        ("Table 3.2", "REST API endpoints (summary)"),
        ("Table 4.1", "Face quality filter thresholds"),
        ("Table 4.2", "Recognition pipeline hyperparameters"),
        ("Table 5.1", "Per-photo recognition results"),
        ("Table 5.2", "Ablation study: impact of pipeline components"),
        ("Table 6.1", "Alignment with UN Sustainable Development Goals"),
    ]
    for num, title in tabs:
        p = doc.add_paragraph()
        run = p.add_run(f"{num}\t")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
        run2 = p.add_run(title)
        run2.font.size = Pt(11)
        run2.font.name = "Times New Roman"
        _set_spacing(p, before=0, after=3, line=1.2)
    add_page_break(doc)


# =============================================================================
# Abstract
# =============================================================================

def build_abstract(doc):
    add_heading_styled(doc, "Abstract", level=1)

    abstract_text = (
        "Manual classroom attendance is slow, error-prone, and susceptible to proxy "
        "responses. This report presents a full-stack Smart Attendance System that "
        "replaces the traditional roll-call workflow with a single classroom "
        "photograph. A teacher uploads the image through a React dashboard; a "
        "Node.js/TypeScript backend authenticates the session and forwards the image to "
        "a Python FastAPI AI service; the service detects every face using MTCNN with a "
        "RetinaFace fallback, filters low-quality crops, generates a 512-d FaceNet "
        "embedding for each face, and passes the embeddings through a lightweight 128-d "
        "projection head trained with triplet loss on the institute's own dataset to "
        "bridge the domain gap between selfie enrollment and uncontrolled classroom "
        "photos. Matching is performed by minimum L2 distance against stored student "
        "encodings, with a support-vector-machine classifier providing an independent "
        "confirmation signal and a margin-based confidence score flagging ambiguous "
        "matches for teacher review. The system enrolls 113 students across two batches "
        "with 4,134 embeddings in total; on 14 held-out classroom group photographs "
        "containing 243 faces it attains an overall recognition rate of 86%, reaching "
        "100% on nine of those photographs while processing each image in one to five "
        "seconds on commodity hardware. The remaining errors cluster on small, "
        "pose-extreme faces in dense groups \u2014 cases where an additional photograph "
        "or simple manual override is a practical fallback. The report documents the "
        "complete engineering trajectory, including two major model pivots "
        "(ArcFace \u2192 FaceNet \u2192 FaceNet + projection head) and the threshold, "
        "margin, and resolution ablations that drove the final configuration."
    )
    add_para(doc, abstract_text, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             after=12, line=1.5)

    add_para(doc, "Keywords: ", size=11, bold=True, after=2)
    add_para(doc,
             "face recognition, attendance system, MTCNN, FaceNet, triplet loss, "
             "projection head, support vector machine, full-stack application, "
             "Node.js, FastAPI, React, Prisma.",
             size=11, italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=10, line=1.3)

    add_page_break(doc)


# =============================================================================
# FYDP as Complex Engineering Problem
# =============================================================================

def build_cep(doc):
    add_heading_styled(doc, "Final Year Design Project as a Complex Engineering Problem", 1)
    add_para(doc,
             "The Washington Accord defines a Complex Engineering Problem (CEP) as one "
             "that cannot be resolved without in-depth engineering knowledge and that "
             "simultaneously satisfies several of the attributes listed below. This "
             "section argues that the Smart Attendance System meets every one of these "
             "attributes and therefore qualifies as a CEP in the sense intended by the "
             "accreditation framework.",
             after=8, line=1.3, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    attrs = [
        ("In-depth engineering knowledge (WK1\u2013WK4):",
         "Solving the project required deep-learning theory (FaceNet, MTCNN, triplet "
         "loss, SVM classification), software-engineering practice (REST API design, "
         "relational schema design, JWT authentication, role-based access control), "
         "and systems engineering (multi-service orchestration, image pipelines, "
         "request tracing, and timeouts). No single undergraduate course covered more "
         "than a slice of this body of knowledge."),

        ("Range of conflicting requirements:",
         "Recognition accuracy, processing latency, and dataset effort pull in opposite "
         "directions. A larger resize preserves detail but slows detection; a tighter "
         "threshold reduces false positives but misses real matches; a deeper projection "
         "head discriminates better on the training set but may overfit a limited "
         "dataset. Every major design decision in Chapters 3\u20134 resolved such a "
         "conflict."),

        ("Depth of analysis required:",
         "The final recognition threshold, margin filter, and resize resolution are the "
         "outputs of controlled ablation studies (Chapter 5), not intuition. The team "
         "compared ArcFace against FaceNet, ran a triplet-loss training on held-out "
         "students, and measured recall vs false-positive rate across eight thresholds "
         "before committing to the production values."),

        ("Familiarity of issues:",
         "Domain-gap problems between controlled enrollment photos and uncontrolled "
         "deployment photos are well known in the face-recognition literature, but the "
         "specific manifestation here \u2014 classroom group photos with dozens of small "
         "faces, variable lighting, and partial occlusion \u2014 demanded a project-"
         "specific solution, namely the enrichment step that adds classroom crops to "
         "the enrollment store."),

        ("Extent of applicable codes:",
         "The system handles personally-identifiable biometric data and must therefore "
         "be designed with ISO/IEC 27001 (information-security management) and ISO "
         "9001:2015 (quality management) in mind, even though no formal audit is "
         "performed. Data-protection obligations under Pakistan\u2019s draft Personal Data "
         "Protection Bill also shape our approach to encoding storage and retention."),

        ("Extent of stakeholder involvement:",
         "Teachers, students, departmental administrators, and IT staff all have "
         "different needs \u2014 teachers want a one-click workflow, students want "
         "privacy and fairness, administrators want auditable logs, and IT staff want "
         "predictable resource usage. The system\u2019s role-based access control and "
         "manual-override UI emerged directly from those conflicting expectations."),

        ("Interdependence:",
         "A change in resize resolution cascades through MTCNN crop sizes, FaceNet "
         "embeddings, projection-head responses, and SVM scores; no component could be "
         "tuned in isolation. This interdependence is characteristic of a CEP."),
    ]

    for head, body in attrs:
        p = doc.add_paragraph()
        r1 = p.add_run(head + " ")
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.name = "Times New Roman"
        r2 = p.add_run(body)
        r2.font.size = Pt(11)
        r2.font.name = "Times New Roman"
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_spacing(p, before=0, after=8, line=1.3)

    add_page_break(doc)


# =============================================================================
# Chapter 1: Introduction
# =============================================================================

def build_chapter_1(doc):
    add_heading_styled(doc, "1. Introduction", level=1)

    # 1.1 Background
    add_heading_styled(doc, "1.1 Background and Motivation", level=2)

    add_para(doc,
             "Taking attendance is one of the oldest classroom rituals: the teacher reads "
             "names, students respond, marks are scribbled in a register, and the register "
             "is eventually copied into a departmental spreadsheet. At GIK Institute \u2014 "
             "a university where many courses enroll between thirty and sixty students \u2014 "
             "this ritual consumes between five and ten minutes of every lecture. Scaled "
             "across a full semester of twenty-four contact sessions, that is roughly four "
             "hours of lost instruction per course per section, or, if every course "
             "followed the same pattern, more than two hundred hours of classroom time "
             "consumed by administrative overhead at the faculty level per semester.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=8)

    add_para(doc,
             "The cost is not only temporal. The manual workflow is also error-prone. "
             "Students who arrive late are often marked absent by mistake; students who "
             "attend occasionally persuade classmates to answer on their behalf (proxy "
             "attendance); registers are lost or transcribed incorrectly; and the "
             "aggregated data available to faculty offices rarely reflects the reality of "
             "day-to-day attendance. Academic rules that depend on attendance \u2014 the "
             "75% attendance threshold for permission to sit final examinations, for "
             "example \u2014 therefore run on data of uneven quality.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=8)

    add_para(doc,
             "Three trends argue that this need not be the case. First, high-resolution "
             "cameras are universal: almost every teacher carries a smartphone capable of "
             "capturing a 3024\u00d74032 pixel image of the classroom. Second, deep-learning "
             "face recognition has matured dramatically. Models such as FaceNet "
             "[Schroff et al., 2015] achieve above 99% accuracy on standard face "
             "verification benchmarks, and open-source detectors such as MTCNN "
             "[Zhang et al., 2016] run on commodity CPUs in seconds. Third, the software "
             "stack needed to wrap these models in a usable web application \u2014 a "
             "JavaScript frontend, a typed backend, a managed database, and a Python AI "
             "service \u2014 is by 2026 a familiar, maintainable set of tools that any "
             "mid-size institute can host on its own infrastructure.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=8)

    add_para(doc,
             "Put these trends together and a different attendance workflow becomes "
             "plausible. A teacher enters the classroom, takes a single photograph, "
             "uploads it through a browser tab, and sees an annotated image in which every "
             "enrolled student's face is identified with a confidence score. The teacher "
             "skims the results, corrects a handful of low-confidence matches, and "
             "submits. Total time: well under one minute. That is the workflow this "
             "project set out to build and to validate with real classroom data.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    # 1.2 Problem Statement
    add_heading_styled(doc, "1.2 Problem Statement", level=2)
    add_para(doc,
             "Manual classroom attendance is slow, error-prone, and vulnerable to proxy "
             "attendance, while modern smartphone cameras and open-source face-recognition "
             "models make a photograph-based alternative technically feasible. The "
             "problem addressed by this project is to design and implement a reliable, "
             "production-grade attendance system that (i) takes a classroom group "
             "photograph as its only per-session input from the teacher, (ii) identifies "
             "every enrolled student in the photograph with a published confidence score, "
             "(iii) allows the teacher to verify and adjust the results before the "
             "session is finalized, and (iv) persists auditable attendance records in a "
             "manner that faculty administrators can inspect and aggregate. The system "
             "must handle the realistic range of classroom photographs \u2014 variable "
             "lighting, small or partially-occluded faces, pose variation \u2014 while "
             "preserving student privacy through role-based access control and local-only "
             "encoding storage.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    # 1.3 Scope
    add_heading_styled(doc, "1.3 Scope of the Work and Expected Outcomes", level=2)
    add_para(doc, "The scope of the project covers:", size=11, line=1.3, after=4)
    add_bullets(doc, [
        "A web-based full-stack application (React frontend, Node.js + Express + "
        "TypeScript backend, Python FastAPI AI service) deployable on a single "
        "Linux server with modest hardware requirements (4 CPU cores, 8 GB RAM).",
        "Role-based access control with two roles: Admin (student, course, and "
        "timetable management) and Teacher (attendance marking and history).",
        "A face-recognition pipeline that attains \u2265 85% recognition rate on the "
        "institute's own classroom photographs.",
        "Enrollment of at least 100 students with 3\u201310 selfie photographs each, "
        "plus enrichment from classroom group photos to bridge the domain gap.",
        "Per-photo processing latency under 10 seconds for images up to 4K resolution.",
        "Auditable attendance records persisted in a relational database, with an "
        "admin interface to browse audit logs.",
    ])

    add_para(doc, "The following are explicitly out of scope:", size=11, line=1.3,
             after=4, before=8)
    add_bullets(doc, [
        "Hardware camera deployment (dedicated classroom cameras, Raspberry-Pi "
        "capture nodes, or embedded IoT devices).",
        "Mobile-native applications; the system is browser-only.",
        "Formal security certification under ISO/IEC 27001 or similar frameworks.",
        "Cloud-scale elasticity (horizontal auto-scaling, multi-region deployment).",
    ])

    add_para(doc, "Expected outcomes at the conclusion of the project are:",
             size=11, line=1.3, after=4, before=8)
    add_numbered(doc, [
        "A running, demonstrable web application with seeded admin and teacher "
        "accounts.",
        "A face-recognition service achieving the accuracy targets outlined above, "
        "validated on held-out classroom photographs.",
        "A reproducible training pipeline for the projection head and SVM classifier, "
        "documented in this report and in the project repository.",
        "A written report (this document) describing the system and its empirical "
        "evaluation.",
    ])

    # 1.4 Outline
    add_heading_styled(doc, "1.4 Outline of Report", level=2)
    add_para(doc,
             "The remainder of this report is organised as follows. Chapter 2 surveys "
             "the face-recognition literature that informed our design choices. Chapter 3 "
             "describes the overall system architecture \u2014 frontend, backend, database, "
             "and AI service \u2014 and the data flow for an attendance session. Chapter 4 "
             "details the recognition pipeline itself, including detection, preprocessing, "
             "embedding, projection-head training, matching, and deduplication. Chapter 5 "
             "reports the empirical evaluation: dataset statistics, ablation studies, and "
             "per-photo results. Chapter 6 discusses impact, sustainability, and "
             "stakeholder considerations. Chapter 7 concludes and outlines future work. "
             "Appendices contain the complete database schema, the REST API reference, and "
             "the hyperparameters used during training.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    # 1.5 Project management / Gantt
    add_heading_styled(doc, "1.5 Project Management and Gantt Chart", level=2)
    add_para(doc,
             "The project spanned twenty-two weeks across the fall 2025 and spring 2026 "
             "semesters. Weekly advisor meetings kept scope aligned, and weekly internal "
             "meetings split work along faculty lines: the two FCSE members handled "
             "backend, frontend, and AI integration while the two FEE members led "
             "dataset collection, enrollment workflow, and hardware/deployment planning. "
             "The Gantt chart in Figure 1.1 shows the high-level timeline.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=6)

    add_figure(doc, "fig_gantt.png",
               "Figure 1.1: Project Gantt chart showing major phases across "
               "September 2025 \u2013 April 2026.", width_inches=6.5)

    # 1.6 Contribution by group members
    add_heading_styled(doc, "1.6 Contribution by Group Members", level=2)
    add_para(doc,
             "While every decision was made collectively, day-to-day responsibilities were "
             "divided as summarized in Table 1.1.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=6)

    add_table(doc,
              headers=["Team Member", "Department", "Primary Responsibilities"],
              rows=[
                  ["Hassaan Ahmed", "FCSE",
                   "Backend (Node.js + Express + Prisma), AI service integration, recognition "
                   "pipeline engineering, database schema design, authentication & RBAC, "
                   "overall system architecture."],
                  ["Osaid Khan Afridi", "FCSE",
                   "Frontend (React + Vite), teacher/admin dashboards, attendance session UI, "
                   "API client, results-review interface, UX polish."],
                  ["Saad Khan", "FEE",
                   "Dataset collection & curation, enrollment workflow, model evaluation, "
                   "projection-head training experiments, testing protocol."],
                  ["Najaf", "FEE",
                   "Hardware/deployment research, networking, documentation, "
                   "user-acceptance testing, Gantt tracking."],
              ],
              caption="Table 1.1: Contribution of team members.",
              col_widths=[1.5, 1.0, 4.0])

    add_page_break(doc)


# =============================================================================
# Chapter 2: Literature Review
# =============================================================================

def build_chapter_2(doc):
    add_heading_styled(doc, "2. Literature Review", level=1)

    add_heading_styled(doc, "2.1 Face Detection", level=2)
    add_para(doc,
             "Modern face detection is a mature field. The Viola-Jones detector [2001] "
             "introduced a cascade of Haar-feature classifiers that still ships with "
             "OpenCV; it is fast but brittle on non-frontal poses and non-uniform "
             "lighting, both of which dominate classroom photographs. Deep-learning-based "
             "detectors supplanted Viola-Jones in the mid-2010s. MTCNN [Zhang et al., "
             "2016] uses a three-stage cascade of convolutional neural networks (P-Net, "
             "R-Net, O-Net) to jointly predict bounding boxes and five facial landmarks. "
             "RetinaFace [Deng et al., 2020] builds on feature-pyramid networks and adds "
             "dense five-point landmarks and a mesh regression branch, giving very high "
             "accuracy on small and profile faces at the cost of roughly 3\u00d7 the MTCNN "
             "inference time.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=8)

    add_para(doc,
             "For our use case \u2014 a single classroom image with dozens of faces and "
             "CPU-only hosting \u2014 we adopted a hybrid strategy: MTCNN as the primary "
             "detector at two scales (full-resolution and 640-px), with RetinaFace invoked "
             "as a secondary detector when it is available in the deployment environment. "
             "Duplicates across the two detectors are merged by non-maximum suppression "
             "with an intersection-over-union threshold of 0.4. This multi-detector "
             "strategy reliably picks up small and partially-occluded faces that either "
             "detector alone would miss.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_heading_styled(doc, "2.2 Face Recognition and Embeddings", level=2)
    add_para(doc,
             "Face recognition is conventionally posed as two problems: verification "
             "(\u201Care these two faces the same person?\u201D) and identification "
             "(\u201Cwho is this face?\u201D). Modern systems solve both by learning an "
             "embedding \u2014 a fixed-length vector in a high-dimensional space \u2014 "
             "such that the embeddings of the same person are close (by Euclidean or "
             "cosine distance) and embeddings of different people are far apart.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=8)

    add_para(doc,
             "FaceNet [Schroff et al., 2015] introduced the triplet-loss formulation and "
             "produced 128- or 512-d embeddings trained on roughly 200 million images. It "
             "reports 99.63% accuracy on the Labeled Faces in the Wild (LFW) benchmark. "
             "The keras_facenet implementation we used exposes a FaceNet model "
             "pre-trained on VGGFace2, producing 512-dimensional embeddings.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=8)

    add_para(doc,
             "ArcFace [Deng et al., 2019] replaces the triplet loss with an angular "
             "additive margin on the softmax classification of training identities. On "
             "benchmarks ArcFace edges out FaceNet (~99.82% on LFW), and the InsightFace "
             "library provides an accessible CPU implementation. Our initial pipeline "
             "used ArcFace, but we observed that the cosine-distance similarity between "
             "a student's selfie enrollment and his or her classroom crop was systematically "
             "lower than between two classroom crops of different students \u2014 a "
             "domain-gap phenomenon more pronounced with ArcFace than with the older, "
             "triplet-trained FaceNet. We therefore migrated to FaceNet after a controlled "
             "comparison (reported in Chapter 5).",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_heading_styled(doc, "2.3 Domain Adaptation for Face Recognition", level=2)
    add_para(doc,
             "A recurring issue in deployed face-recognition systems is the domain gap "
             "between enrollment photographs (usually frontal, well-lit, close-cropped "
             "selfies) and recognition photographs (ambient lighting, off-axis poses, "
             "small crops). The standard palliatives are either (a) enrolling students "
             "from in-domain photographs when possible, or (b) training a lightweight "
             "adapter on a small in-domain dataset. Projection heads trained with triplet "
             "loss \u2014 popularized in self-supervised learning [Chen et al., 2020, "
             "SimCLR] and since applied to face recognition by several industrial teams "
             "\u2014 take a pretrained backbone embedding (here, FaceNet's 512-d vector) "
             "and learn a small feed-forward network that produces a more discriminative "
             "in-domain embedding.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=8)

    add_para(doc,
             "We adopted this approach: a 512 \u2192 256 \u2192 128 dense network with "
             "ReLU activations, batch normalisation, and a final L2-normalisation layer, "
             "trained with a triplet margin of 0.3 for up to eighty epochs with early "
             "stopping. The projection head training dataset is built from the same "
             "enrollment encodings the system already stores; the positive pairs come from "
             "the same student's multiple selfies and enriched classroom crops, and the "
             "negatives are sampled uniformly across the other 112 students.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_heading_styled(doc, "2.4 Classifier Layer: SVMs for Face Identification", level=2)
    add_para(doc,
             "Once an embedding space is learned, identification can be performed by "
             "nearest-neighbour search (minimum Euclidean distance) or by a learned "
             "classifier. Radial-basis-function SVMs were a standard choice in the pre-"
             "deep-learning era and remain competitive as a confirmation layer: they "
             "produce a calibrated probability for each class, they are cheap to retrain "
             "when new students enroll, and they give independent evidence that does not "
             "fully correlate with the nearest-neighbour distance. Our system uses an "
             "RBF-kernel SVM with hyperparameters chosen by three-fold grid search "
             "(C \u2208 {1, 10, 50, 100}, gamma \u2208 {scale, auto}) trained on all "
             "4,134 projected encodings.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_heading_styled(doc, "2.5 Comparable Systems", level=2)
    add_para(doc,
             "A number of commercial and academic systems address classroom attendance "
             "with face recognition. Products such as NEC NeoFace, Megvii Face++, and "
             "Amazon Rekognition offer turnkey cloud APIs; however, they carry "
             "per-image or subscription fees that make them impractical for a university "
             "with dozens of parallel sections, and they require sending biometric data "
             "off-premises. Academic work \u2014 for example Lamba et al. [2021] and "
             "Aravind et al. [2019] \u2014 reports systems with single-digit-student "
             "datasets and rarely tackles the dense-classroom case of thirty to forty "
             "simultaneous faces. Our system positions itself in this gap: on-premises "
             "deployment, full-stack workflow, institute-scale dataset, and an explicit "
             "focus on uncontrolled classroom photographs.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=8)

    # Comparison table
    add_table(doc,
              headers=["Backbone", "Output dim", "LFW Acc.", "In-Domain Behaviour (our tests)"],
              rows=[
                  ["ArcFace (iResNet-100)", "512", "~99.82%",
                   "Strong on frontal selfies; poorer generalization from selfie to "
                   "classroom crop on our dataset."],
                  ["FaceNet (Inception-v1)", "512", "~99.63%",
                   "More stable across selfie/classroom gap; chosen as production backbone."],
                  ["FaceNet + projection head", "128", "n/a",
                   "Best intra-/inter-class separation on our dataset; tighter threshold "
                   "possible; final production model."],
              ],
              caption="Table 2.1: Comparison of surveyed face-recognition backbones.",
              col_widths=[1.8, 1.0, 1.2, 2.5])

    add_heading_styled(doc, "2.6 Reasoning Taken from the Literature", level=2)
    add_para(doc,
             "From the literature we drew four concrete design decisions: (i) use a "
             "multi-detector cascade rather than a single detector, to tolerate small and "
             "profile faces; (ii) prefer FaceNet over ArcFace for our cross-domain "
             "setting despite ArcFace\u2019s higher benchmark numbers; (iii) add a small "
             "projection head trained in-domain to bridge the selfie/classroom gap; and "
             "(iv) keep an SVM classifier as a second opinion that does not share the "
             "nearest-neighbour failure modes.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_page_break(doc)


# =============================================================================
# Chapter 3: System Design and Architecture
# =============================================================================

def build_chapter_3(doc):
    add_heading_styled(doc, "3. System Design and Architecture", level=1)

    add_heading_styled(doc, "3.1 Overview", level=2)
    add_para(doc,
             "The Smart Attendance System follows a classical three-tier architecture "
             "extended with a dedicated AI service. Figure 3.1 shows the high-level "
             "components and their interactions. The frontend is a React + Vite single-"
             "page application served as static assets; the backend is a Node.js + "
             "Express + TypeScript REST API that owns all persistent state through a "
             "Prisma-managed SQLite database; the AI service is a Python + FastAPI "
             "process that wraps the face-detection and face-recognition models. The "
             "three processes communicate over HTTP and are designed to be deployable on "
             "a single Linux host (our target is an Oracle Cloud Ampere A1 VM with 4 OCPU "
             "and 24 GB RAM, offered under Oracle\u2019s always-free tier).",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=8)

    add_figure(doc, "fig_architecture.png",
               "Figure 3.1: High-level system architecture.", width_inches=6.5)

    add_heading_styled(doc, "3.2 Frontend", level=2)
    add_para(doc,
             "The frontend is built with React 18 and Vite 7; TypeScript is used for "
             "type safety. Routing is handled by React Router; state management follows "
             "a lightweight Context + hooks pattern with no external state library. The "
             "UI is divided into three role-specific dashboards: an Admin dashboard "
             "(student, teacher, course, classroom, and timetable CRUD; audit-log "
             "browser), a Teacher dashboard (schedule view, attendance session flow, "
             "attendance history), and a Testing dashboard used during development to "
             "call the recognition endpoint directly against arbitrary photographs.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=8)

    add_para(doc,
             "The attendance session flow is the frontend\u2019s most complex feature. The "
             "teacher selects a course, a date, and a time slot; the frontend creates a "
             "session record via the backend API and then presents an image-upload control. "
             "When the teacher submits a photograph, the frontend transparently forwards "
             "it to the backend and displays a loading indicator. On receipt of the "
             "recognition result, the UI renders the annotated image with bounding boxes, "
             "each labelled with the recognized student\u2019s name and a coloured "
             "confidence pill (green \u2265 50%, yellow 15\u201349%, red < 15%). Below the "
             "image a reviewable list allows the teacher to flip any match from "
             "\u201CPresent\u201D to \u201CAbsent\u201D, override a misrecognition, or add "
             "a student who was missed. Only after the teacher clicks \u201CFinalize\u201D "
             "are attendance records persisted.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_heading_styled(doc, "3.3 Backend", level=2)
    add_para(doc,
             "The backend is a standard Express application organised into feature "
             "modules (auth, students, teachers, courses, classrooms, timetables, "
             "attendance, audit, testing). Each module follows a strict four-file "
             "convention: service.ts holds business logic and the Prisma queries, "
             "controller.ts is the thin HTTP layer, validation.ts contains Zod-style "
             "request validation, and routes.ts registers the endpoints. This "
             "convention is mechanical enough that new endpoints can be added without "
             "reading more than the file that owns them, which paid dividends during "
             "the accelerated final weeks of implementation.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=8)

    add_para(doc,
             "Authentication is based on JSON Web Tokens. On login, the backend "
             "returns a short-lived access token (15 minutes) and a long-lived refresh "
             "token (seven days). Refresh tokens are stored hashed in the database and "
             "a revocation table allows us to invalidate sessions. Role-based access "
             "control is enforced through two lightweight middleware functions, "
             "`adminOnly` and `teacherOnly`, which inspect the decoded JWT payload; "
             "many endpoints are wrapped in both to permit either role.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_heading_styled(doc, "3.4 Database Schema", level=2)
    add_para(doc,
             "Persistent state lives in a Prisma-managed SQLite database. SQLite was "
             "chosen for ease of deployment (single-file database, no separate server "
             "process) and for its performance characteristics at the scale we target "
             "(thousands of students, tens of thousands of attendance records). Prisma "
             "provides a type-safe query builder and automatic migration tooling. The "
             "core entities are listed in Table 3.1, and their relationships are shown "
             "in Figure 3.2.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=8)

    add_table(doc,
              headers=["Entity", "Purpose"],
              rows=[
                  ["User", "Admin and teacher accounts; owns authentication fields."],
                  ["Teacher", "Profile extension of User for teachers (employee ID, office, etc.)."],
                  ["Student", "Enrolled students; primary key is registration number."],
                  ["Department, Batch",
                   "Reference tables for organisational grouping."],
                  ["Course, TeacherCourseAssignment",
                   "Course catalogue and which teacher teaches what."],
                  ["Classroom, TimetableSlot",
                   "Physical classrooms and the weekly timetable."],
                  ["StudentEnrollment",
                   "M:N mapping of students to courses for a given term."],
                  ["AttendanceSession, AttendanceRecord",
                   "A session is a single lecture instance; a record is one student\u2019s "
                   "status for that session."],
                  ["FaceEncoding", "Metadata; the actual NumPy arrays live on disk."],
                  ["AuditLog", "Append-only log of every sensitive operation."],
                  ["RefreshToken", "Hashed refresh tokens for JWT rotation/revocation."],
              ],
              caption="Table 3.1: Core database entities.",
              col_widths=[2.0, 4.5])

    add_figure(doc, "fig_erd.png",
               "Figure 3.2: Simplified entity-relationship diagram.", width_inches=6.5)

    add_heading_styled(doc, "3.5 AI Service", level=2)
    add_para(doc,
             "The AI service is a FastAPI application with three primary endpoints: "
             "`/api/v1/recognize` (process a classroom photo), `/api/v1/encodings` "
             "(enroll or delete a student), and `/api/v1/health` (readiness probe). "
             "Models are loaded once into memory at process start through a singleton "
             "ModelLoader class, which at present loads MTCNN, RetinaFace (optional), "
             "FaceNet, and the custom projection head. Requests are serviced in a "
             "thread-pool executor so that the asyncio event loop remains responsive "
             "under concurrent load.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=8)

    add_para(doc,
             "The service is deliberately stateless: it does not know about courses, "
             "sessions, or teachers. It accepts a JPEG/PNG image and a JSON array of "
             "student IDs, and it returns a list of recognized students (with bounding "
             "boxes, names, and confidences) and a list of unknown faces. Keeping the AI "
             "service stateless means it can be scaled, restarted, and replaced "
             "independently of the backend.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_heading_styled(doc, "3.6 Request Flow for an Attendance Session", level=2)
    add_para(doc,
             "Figure 3.3 traces an attendance session from the teacher\u2019s first click "
             "to the finalized attendance records. The frontend never talks to the AI "
             "service directly; the backend brokers every call, enforces authentication, "
             "and attaches the list of enrolled student IDs to the recognition request so "
             "that the AI service constrains its search to students actually registered "
             "for the course.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=8)

    add_figure(doc, "fig_dataflow.png",
               "Figure 3.3: Attendance session data flow.", width_inches=6.5)

    add_heading_styled(doc, "3.7 REST API Surface", level=2)
    add_para(doc,
             "Table 3.2 summarises the REST endpoints; full request/response shapes are "
             "given in Appendix B.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=6)

    add_table(doc,
              headers=["Group", "Endpoint", "Role"],
              rows=[
                  ["Auth", "POST /auth/login, /refresh, /logout", "Any"],
                  ["Auth", "GET /auth/profile", "Any"],
                  ["Admin", "GET/POST/PUT/DELETE /admin/students", "Admin"],
                  ["Admin", "GET/POST/PUT/DELETE /admin/teachers", "Admin"],
                  ["Admin", "GET/POST/PUT/DELETE /admin/courses", "Admin"],
                  ["Admin", "GET/POST/PUT/DELETE /admin/classrooms", "Admin"],
                  ["Admin", "GET/POST/PUT/DELETE /admin/timetables", "Admin"],
                  ["Admin", "GET /admin/logs/audit, /attendance", "Admin"],
                  ["Teacher", "GET /teacher/schedule", "Teacher"],
                  ["Teacher", "POST /teacher/sessions", "Teacher"],
                  ["Teacher", "POST /teacher/sessions/:id/image", "Teacher"],
                  ["Teacher", "GET /teacher/sessions/:id", "Teacher"],
                  ["Teacher", "POST /teacher/sessions/:id/finalize", "Teacher"],
                  ["AI", "POST /api/v1/recognize", "internal"],
                  ["AI", "POST /api/v1/encodings", "internal"],
              ],
              caption="Table 3.2: Summary of REST API endpoints.",
              col_widths=[1.0, 3.5, 1.3])

    add_heading_styled(doc, "3.8 Summary", level=2)
    add_para(doc,
             "The architecture cleanly separates concerns: the frontend is presentation-"
             "only, the backend is the system of record, and the AI service is a "
             "stateless, horizontally-replaceable inference engine. This separation made "
             "it possible to replace the recognition backbone twice (ArcFace \u2192 "
             "FaceNet \u2192 FaceNet + projection head) without touching the frontend or "
             "the database.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_page_break(doc)


# =============================================================================
# Chapter 4: Face Recognition Pipeline
# =============================================================================

def build_chapter_4(doc):
    add_heading_styled(doc, "4. Face Recognition Pipeline", level=1)

    add_heading_styled(doc, "4.1 Pipeline Overview", level=2)
    add_para(doc,
             "Figure 4.1 shows the seven-stage recognition pipeline. Each stage is "
             "described in detail in the subsections below.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=6)

    add_figure(doc, "fig_pipeline.png",
               "Figure 4.1: End-to-end face recognition pipeline.", width_inches=6.5)

    add_heading_styled(doc, "4.2 Image Ingestion and Resize", level=2)
    add_para(doc,
             "Incoming classroom photographs range from 1920\u00d71080 pixel phone snaps to "
             "5712\u00d74284 pixel DSLR photos. Processing a 5712-pixel image directly "
             "through MTCNN takes over sixty seconds on our target hardware \u2014 "
             "unacceptable for an interactive workflow. We therefore resize to a maximum "
             "dimension of 2048 pixels before detection, which reduces processing time to "
             "under five seconds while preserving sufficient detail for reliable "
             "recognition of faces down to about twenty pixels square. The resize uses "
             "OpenCV\u2019s INTER_AREA interpolation for best downsampling quality.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_heading_styled(doc, "4.3 Face Detection", level=2)
    add_para(doc,
             "Faces are detected by two complementary detectors working in parallel: "
             "MTCNN at the resized resolution, plus MTCNN at a further-reduced 640-pixel "
             "resolution to pick up tiny faces, plus RetinaFace when it is available. The "
             "union of detections is fed through non-maximum suppression with an "
             "intersection-over-union threshold of 0.4. A detection is kept if its "
             "bounding box is at least 10 pixels in each dimension and its confidence is "
             "at least 0.90.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_heading_styled(doc, "4.4 Face Preprocessing and Quality Assessment", level=2)
    add_para(doc,
             "Each detected face is cropped exactly as MTCNN reports its bounding box "
             "(no alignment, no margin) \u2014 a deliberately simple choice that "
             "matched the approach used during enrollment and during training of the "
             "projection head. The cropped RGB patch is then passed through a quality "
             "filter (Table 4.1) that rejects tiny, blurry, over-exposed, or extreme-"
             "profile faces before embedding.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=6)

    add_table(doc,
              headers=["Check", "Criterion", "Purpose"],
              rows=[
                  ["Minimum size", "w and h \u2265 20 px",
                   "Avoid tiny crops that embed to noise."],
                  ["Aspect ratio", "0.3 \u2264 w/h \u2264 3.0",
                   "Reject obviously malformed boxes."],
                  ["Side profile", "|eye_x_dist| / w \u2265 0.15",
                   "Reject extreme profile views."],
                  ["Sharpness", "Laplacian variance \u2265 5.0",
                   "Reject motion-blurred crops."],
                  ["Exposure", "15 \u2264 mean_grey \u2264 250",
                   "Reject severely under- or over-exposed crops."],
              ],
              caption="Table 4.1: Face quality filter thresholds.",
              col_widths=[1.5, 2.0, 3.0])

    add_heading_styled(doc, "4.5 FaceNet Embedding", level=2)
    add_para(doc,
             "Each accepted crop is passed to the keras_facenet FaceNet model, which "
             "internally resizes the image to 160\u00d7160 pixels and produces a 512-"
             "dimensional, L2-normalised embedding. The choice of keras_facenet over "
             "alternatives (InsightFace, face_recognition) was motivated by its minimal "
             "dependency footprint, proven behaviour on this dataset (taken from our team "
             "member\u2019s prior notebook `Mtcnn_Final.ipynb`), and its CPU-friendly "
             "inference of roughly 50 ms per face.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_heading_styled(doc, "4.6 Projection Head", level=2)
    add_para(doc,
             "The 512-d FaceNet embedding is then passed through a learned 128-d "
             "projection head. The architecture is a simple three-layer dense network "
             "(512 \u2192 256 \u2192 128) with ReLU activations, batch normalisation, a "
             "10% dropout, and a final L2-normalisation layer. The head is trained with "
             "triplet loss on anchor/positive/negative triplets mined from the enrollment "
             "encodings. Formally, given an anchor embedding \u27e8a, p, n\u27e9 where p "
             "shares the same student label as a and n does not, the loss is:",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=8)

    add_para(doc, "L = max(0, \u2016f(a) \u2212 f(p)\u2016\u00b2 \u2212 \u2016f(a) \u2212 f(n)\u2016\u00b2 + \u03b1)",
             size=11, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, after=8)

    add_para(doc,
             "where f(\u00b7) is the projection head and \u03b1 = 0.3 is the margin. "
             "Training runs for up to eighty epochs with fifty mini-batches of 64 triplets "
             "per epoch; early stopping monitors the held-out validation loss with "
             "patience 15. Figure 4.2 shows the typical training behaviour \u2014 "
             "converging in roughly 60 epochs.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=6)

    add_figure(doc, "fig_training.png",
               "Figure 4.2: Projection head training/validation loss.", width_inches=5.5)

    add_para(doc,
             "After training, the projection head is used both online (projecting query "
             "embeddings at recognition time) and offline (projecting all stored "
             "enrollment encodings into the same 128-d space, via the "
             "`reproject_encodings.py` script). Matching is then performed in the "
             "projected space.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=6)

    add_figure(doc, "fig_separation.png",
               "Figure 4.3: Intra-class vs inter-class L2 distance distribution in "
               "the projected 128-d space. The chosen threshold t = 1.1 is marked.",
               width_inches=5.5)

    add_heading_styled(doc, "4.7 Matching Strategy", level=2)
    add_para(doc,
             "For each query face we compute the L2 distance to every stored encoding of "
             "every enrolled student (constrained to the student IDs the backend "
             "provided). The best match is the student with the smallest minimum distance. "
             "The match is accepted if that smallest distance is below a global threshold "
             "(t = 1.1 in production) and is converted into a confidence score:",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=6)

    add_para(doc, "conf_raw = max(0, 1 \u2212 d_best / t) \u00d7 100",
             size=11, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, after=6)

    add_para(doc,
             "A margin penalty reduces the raw confidence when the best match is not much "
             "closer than the second-best, because this indicates ambiguity between "
             "students who look similar:",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=6)

    add_table(doc,
              headers=["Margin (d\u2082 \u2212 d\u2081)", "Confidence multiplier"],
              rows=[
                  ["< 0.03", "\u00d7 0.4"],
                  ["0.03 \u2013 0.06", "\u00d7 0.6"],
                  ["0.06 \u2013 0.10", "\u00d7 0.75"],
                  ["0.10 \u2013 0.15", "\u00d7 0.85"],
                  ["\u2265 0.15", "\u00d7 1.0 (unchanged)"],
              ],
              caption="",
              col_widths=[2.5, 2.5])

    add_para(doc,
             "A separate SVM classifier (trained on the same projected encodings) runs in "
             "parallel and is used for confirmation only: if the SVM agrees with the "
             "nearest-neighbour result with probability at least 0.30, we blend the SVM "
             "probability into the confidence score via 0.7 \u00b7 conf_raw + "
             "0.3 \u00b7 (100 \u00b7 p_svm) and tag the match as `l2_proj+svm`. Matches "
             "with final confidence below 5% (8% for rematches) are rejected to "
             "\u201CUnknown\u201D.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_heading_styled(doc, "4.8 Deduplication", level=2)
    add_para(doc,
             "A persistent failure mode of any nearest-neighbour matching system is that "
             "two different faces in the same photograph can claim the same identity \u2014 "
             "particularly plausible when a student appears twice because of a reflection, "
             "or when two siblings have similar features. We handle this through a "
             "deduplication pass: after the initial pass, any student claimed by more than "
             "one face keeps only the face with the smallest distance; the other faces "
             "are re-matched against the set of remaining unclaimed students. This "
             "continues for up to ten rounds or until no conflicts remain. Rematched "
             "faces are tagged in the output so the teacher-review UI can surface them.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_heading_styled(doc, "4.9 Hyperparameter Summary", level=2)
    add_para(doc,
             "Table 4.2 summarises the full set of production hyperparameters.",
             size=11, after=6)

    add_table(doc,
              headers=["Parameter", "Value", "Notes"],
              rows=[
                  ["Resize MAX_DIM", "2048 px", "balances speed vs face detail"],
                  ["MTCNN min_face_size", "20 px", "library default"],
                  ["Detection confidence", "\u2265 0.90", "filters spurious boxes"],
                  ["NMS IoU", "0.4", "merges across MTCNN + RetinaFace"],
                  ["Embedding dim (raw)", "512 (FaceNet)", ""],
                  ["Embedding dim (projected)", "128", "\u2193 head: 512\u2192256\u2192128"],
                  ["Projection head triplet margin", "0.3", ""],
                  ["Projection head training epochs", "\u2264 80, early-stop patience 15", ""],
                  ["L2 threshold", "1.1", "projected space"],
                  ["Minimum margin", "0.03 (hard reject below); soft penalty up to 0.15", ""],
                  ["Minimum confidence", "5% (8% rematch)", ""],
                  ["SVM kernel / C / gamma", "RBF / grid-search / grid-search", ""],
                  ["SVM probability threshold", "0.30", "triggers confirmation blend"],
                  ["Dedup rounds", "10", "re-match second-best candidates"],
              ],
              caption="Table 4.2: Recognition pipeline hyperparameters.",
              col_widths=[2.3, 2.0, 2.2])

    add_page_break(doc)


# =============================================================================
# Chapter 5: Results and Discussion
# =============================================================================

def build_chapter_5(doc):
    add_heading_styled(doc, "5. Results and Discussion", level=1)

    add_heading_styled(doc, "5.1 Dataset", level=2)
    add_para(doc,
             "The dataset used throughout this work consists of selfie enrollment "
             "photographs and classroom group photographs from two cohorts of GIK "
             "Institute students: the 2022 batch (sixty-five students) and the 2023 "
             "batch (forty-nine students), for a total of 114 students. After quality "
             "filtering and outlier removal, 113 students have usable encodings in the "
             "store. The selfie dataset contains between three and six photographs per "
             "student, typically captured by the student themselves on their own phone "
             "in ordinary lighting. The classroom test set comprises fourteen group "
             "photographs of varying size, angle, and lighting captured during regular "
             "lectures across the Faculty of Electrical Engineering and the Faculty of "
             "Computer Science and Engineering.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=8)

    add_para(doc,
             "After initial enrollment, the `enrich_from_groups.py` script adds high-"
             "confidence classroom crops to each student\u2019s encoding store, yielding a "
             "final total of 4,134 projected encodings across 113 students (an average of "
             "36.6 encodings per student, with a minimum of 14 and a maximum of 45). "
             "Outlier removal is performed by computing cosine distance from the student "
             "centroid and dropping encodings beyond 2.0 standard deviations; this "
             "eliminated approximately 7% of raw embeddings as likely misdetections or "
             "erroneous enrollment captures.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_heading_styled(doc, "5.2 Model Evolution", level=2)
    add_para(doc,
             "Figure 5.2 summarises how the recognition accuracy evolved across three "
             "major model revisions. The first iteration used ArcFace embeddings with "
             "cosine similarity; on our test set it reached only 55% end-to-end "
             "recognition rate, with observable mis-recognitions (wrong identities "
             "assigned, not just unknowns). Switching to FaceNet raised end-to-end "
             "recognition to 78% and eliminated mis-recognitions entirely, though at the "
             "cost of many \u201Cunknown\u201D faces. Adding the projection head raised "
             "accuracy to 86% while keeping the SVM-confirmed fraction of "
             "mis-recognitions near zero.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=6)

    add_figure(doc, "fig_model_comparison.png",
               "Figure 5.2: Recognition accuracy across model revisions.", width_inches=5.5)

    add_heading_styled(doc, "5.3 Threshold Ablation", level=2)
    add_para(doc,
             "We searched the L2 threshold in the projected space across eight values "
             "(0.80, 0.90, 1.00, 1.05, 1.10, 1.15, 1.20, 1.25). Recall rises sharply up "
             "to about t = 1.1 and then plateaus; the estimated false-positive rate "
             "climbs slowly below t = 1.1 and then accelerates rapidly. We selected "
             "t = 1.1 as the knee of the curve \u2014 it captures 86% recall while "
             "keeping false positives bounded by the margin filter.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=6)

    add_figure(doc, "fig_threshold.png",
               "Figure 5.1: Threshold ablation in the projected 128-d space.",
               width_inches=5.8)

    add_heading_styled(doc, "5.4 Per-Photo Results", level=2)
    add_para(doc,
             "Table 5.1 and Figure 5.3 report the full per-photo breakdown across the "
             "fourteen test photographs, totalling 243 detected faces. Nine out of "
             "fourteen photographs achieve 92\u2013100% recognition rates; the remaining "
             "five are large classroom groups with many small faces. Across all photos, "
             "209 of 243 faces are recognized (86%), with 34 unknowns and 49 \u201Crisky\u201D "
             "matches that are flagged to the teacher-review UI with a confidence below "
             "15%. No false positives were observed under this configuration \u2014 "
             "every recognized student was actually present in the photograph, verified "
             "manually.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=6)

    add_table(doc,
              headers=["Photo", "Detected", "Recognized", "Unknown", "Low-conf", "Rate"],
              rows=[
                  ["IMG_6939", "1", "1", "0", "0", "100%"],
                  ["IMG_6941", "14", "14", "0", "5", "100%"],
                  ["IMG_6945", "5", "4", "1", "0", "80%"],
                  ["IMG_6948", "11", "11", "0", "0", "100%"],
                  ["IMG_6991", "19", "19", "0", "1", "100%"],
                  ["IMG_6996", "31", "24", "7", "9", "77%"],
                  ["IMG_7042", "34", "20", "14", "7", "59%"],
                  ["IMG_7574", "17", "17", "0", "5", "100%"],
                  ["IMG_7647", "7", "7", "0", "0", "100%"],
                  ["IMG_7649", "23", "22", "1", "2", "96%"],
                  ["IMG_6788", "23", "19", "4", "7", "83%"],
                  ["IMG_6794", "13", "12", "1", "2", "92%"],
                  ["IMG_6796", "23", "18", "5", "6", "78%"],
                  ["IMG_whatsapp", "22", "21", "1", "5", "95%"],
                  ["Total", "243", "209", "34", "49", "86%"],
              ],
              caption="Table 5.1: Per-photo recognition results.",
              col_widths=[1.3, 0.9, 1.0, 0.9, 1.0, 0.8])

    add_figure(doc, "fig_accuracy.png",
               "Figure 5.3: Recognition performance across test photographs.",
               width_inches=6.5)

    add_heading_styled(doc, "5.5 Ablation Study", level=2)
    add_para(doc,
             "Table 5.2 quantifies the incremental contribution of the pipeline\u2019s "
             "major components. Each row removes one component and re-evaluates on the "
             "full test set.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=6)

    add_table(doc,
              headers=["Configuration", "Recognition Rate", "Observation"],
              rows=[
                  ["Full pipeline (production)", "86.0%", "Baseline for comparison."],
                  ["\u2212 projection head (raw 512-d)", "\u2248 78%",
                   "Domain gap reduces cross-source matching."],
                  ["\u2212 SVM confirmation", "\u2248 85%",
                   "SVM helps only ~1% of matches, mostly borderline."],
                  ["\u2212 margin filter", "\u2248 90% raw recall, \u003c but with false positives",
                   "Accepts ambiguous low-margin matches."],
                  ["\u2212 RetinaFace (MTCNN only)", "\u2248 83%",
                   "Small / profile faces missed."],
                  ["\u2212 multi-scale detection", "\u2248 79%",
                   "Tiny faces in large groups missed."],
                  ["\u2212 quality filter", "\u2248 86% but noisier embeddings",
                   "Low-quality crops embed unreliably."],
              ],
              caption="Table 5.2: Ablation study \u2014 impact of individual "
                      "pipeline components.",
              col_widths=[2.5, 2.0, 2.2])

    add_heading_styled(doc, "5.6 Processing Time", level=2)
    add_para(doc,
             "End-to-end processing time on our development laptop (Intel Core i7 "
             "12-core, 16 GB RAM, no GPU) ranged from 1.3 seconds for a single-face "
             "photograph to 4.9 seconds for a 31-face classroom group. The breakdown is "
             "dominated by face detection (approximately 70% of wall time), followed by "
             "embedding (20%) and matching (10%). Because detection time scales with "
             "image resolution rather than face count, the resize to 2048 pixels is the "
             "single most important optimization. GPU acceleration would likely halve "
             "these numbers but is not required for the target workflow.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_heading_styled(doc, "5.7 Failure Analysis", level=2)
    add_para(doc,
             "A qualitative inspection of the 34 unknown faces revealed three failure "
             "modes: (i) very small faces (fewer than 25 pixels on the shorter side), "
             "which together with motion blur produce embeddings far from any stored "
             "encoding; (ii) extreme pose variation (profile or near-profile views), "
             "which the quality filter flags as side profiles and skips; and (iii) "
             "students whose enrollment data is sparse (fewer than 20 encodings), where "
             "the projection head has fewer positive pairs to generalise from. Of the "
             "49 low-confidence matches, manual inspection suggested that roughly 35 "
             "were correct, 10 were ambiguous (between two students with similar "
             "features), and 4 were clear mistakes. The teacher-review UI is designed "
             "to surface exactly these 49 matches so that mistakes can be corrected "
             "before finalisation.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_heading_styled(doc, "5.8 Discussion", level=2)
    add_para(doc,
             "Three broader lessons emerged from these experiments. First, benchmark "
             "numbers on curated datasets (LFW, MegaFace) over-predict real-world "
             "accuracy when there is any domain gap between enrollment and deployment "
             "photos \u2014 ArcFace\u2019s higher LFW number did not translate into "
             "better classroom performance on our dataset. Second, simple calibrated "
             "signals (margin, SVM probability) give the system a grammar for flagging "
             "its own uncertainty; without them, the choice is between a permissive "
             "threshold that admits false positives and a strict threshold that admits "
             "many unknowns. Third, the system benefits from domain-specific enrichment "
             "at least as much as from model upgrades: adding classroom crops to each "
             "student\u2019s encoding store raised recognition by more than the switch "
             "from ArcFace to FaceNet.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_page_break(doc)


# =============================================================================
# Chapter 6: Impact & Economic Analysis
# =============================================================================

def build_chapter_6(doc):
    add_heading_styled(doc, "6. Impact and Economic Analysis", level=1)

    add_heading_styled(doc, "6.1 Sustainable Development Goals", level=2)
    add_para(doc,
             "The system contributes to three of the United Nations Sustainable "
             "Development Goals (Table 6.1):",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=6)

    add_table(doc,
              headers=["SDG", "Goal", "Contribution"],
              rows=[
                  ["SDG 4", "Quality Education",
                   "Reclaims classroom time lost to roll call; more reliable "
                   "attendance data improves pedagogical decision-making and academic "
                   "interventions."],
                  ["SDG 9", "Industry, Innovation & Infrastructure",
                   "Demonstrates a deployable, on-premises AI system in an academic "
                   "institution; the full-stack template can be re-used for other "
                   "campus automation."],
                  ["SDG 16", "Peace, Justice & Strong Institutions",
                   "Auditable attendance records reduce opportunities for proxy "
                   "attendance and administrative manipulation of records."],
              ],
              caption="Table 6.1: Alignment with UN Sustainable Development Goals.",
              col_widths=[0.8, 2.0, 3.7])

    add_heading_styled(doc, "6.2 Social Impact", level=2)
    add_para(doc,
             "Student privacy is the dominant social consideration. The system stores "
             "only derived numerical embeddings \u2014 not the original face "
             "photographs \u2014 and does so locally in the institute\u2019s own storage; "
             "no biometric data leaves the deployment host. Enrollment is opt-in and "
             "revocable: on graduation or on explicit request, a student\u2019s encoding "
             "directory is deleted and no trace of the underlying face remains in the "
             "system. Role-based access control prevents teachers from viewing "
             "attendance data that does not concern their own sections.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_heading_styled(doc, "6.3 Economic Analysis", level=2)
    add_para(doc,
             "The system is designed for zero recurring cost to the institute. The "
             "reference deployment targets Oracle Cloud\u2019s always-free Ampere A1 VM "
             "(4 OCPU, 24 GB RAM, 200 GB disk), which is provisioned at no charge for "
             "unlimited time. Development costs comprised only the students\u2019 time "
             "and minor hardware upgrades; no software licences were paid. Savings are "
             "driven almost entirely by reclaimed classroom time, which at conservative "
             "back-of-envelope valuations (five minutes per lecture, twenty-four lectures "
             "per course per semester, dozens of courses running in parallel at the "
             "faculty level) translate to several hundred person-hours of instruction "
             "per semester.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_heading_styled(doc, "6.4 Sustainability Analysis", level=2)
    add_para(doc,
             "From an environmental standpoint, the system replaces paper registers with "
             "a single low-power VM process; the embedded carbon footprint of a single "
             "shared server hosting an institute-wide application is negligible compared "
             "to the paper, transport, and archival costs of traditional records. The "
             "software stack (Node.js, Python, React) is maintained by active open-source "
             "communities, ensuring that the system will remain patchable and deployable "
             "for the foreseeable future without lock-in to a commercial vendor.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_heading_styled(doc, "6.5 Hazard Identification and Safety Measures", level=2)
    add_para(doc,
             "Being a purely software system, there are no mechanical, electrical, or "
             "chemical hazards. The safety considerations are therefore informational:",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=6)
    add_bullets(doc, [
        "Bias: face-recognition accuracy can vary with skin tone, age, or lighting. "
        "The teacher-review UI is the primary mitigation: no attendance record is "
        "final without a human approval.",
        "Privacy: biometric embeddings are kept on-premises; backups are encrypted "
        "at rest; admin accounts require strong passwords and are audit-logged.",
        "Availability: a down AI service does not block attendance marking \u2014 the "
        "backend gracefully falls back to manual attendance entry and flags the "
        "session for later processing.",
        "Misuse: RBAC prevents teachers from modifying records for sections they do "
        "not teach; all modifications are logged with actor and timestamp.",
    ])

    add_heading_styled(doc, "6.6 Standards", level=2)
    add_para(doc,
             "The system design references two ISO standards even though no formal "
             "audit has been performed:",
             size=11, after=6)
    add_bullets(doc, [
        "ISO 9001:2015 (Quality Management Systems): the backend\u2019s module "
        "convention, code review practice, and logging discipline map directly onto "
        "the quality-management clauses of this standard.",
        "ISO/IEC 27001 (Information Security Management): password hashing with "
        "bcrypt, hashed refresh-token storage, and append-only audit logs implement "
        "the minimum technical controls of this standard.",
    ])

    add_heading_styled(doc, "6.7 Summary", level=2)
    add_para(doc,
             "The Smart Attendance System is an intentionally privacy-conscious, on-"
             "premises, open-source-only system that addresses a real pain point at the "
             "institute with zero recurring cost. Its principal risks are algorithmic "
             "(recognition error) rather than physical, and those risks are bounded by "
             "a mandatory human-in-the-loop review step.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_page_break(doc)


# =============================================================================
# Chapter 7: Conclusion
# =============================================================================

def build_chapter_7(doc):
    add_heading_styled(doc, "7. Conclusion and Future Recommendations", level=1)

    add_heading_styled(doc, "7.1 Conclusion", level=2)
    add_para(doc,
             "This project set out to replace manual classroom attendance with a single "
             "classroom photograph, processed end-to-end by a full-stack system that the "
             "institute could host on its own hardware. Across twenty-two weeks of "
             "development we delivered a React + Node.js + Python web application, a "
             "three-stage face-recognition pipeline (MTCNN + FaceNet + custom 128-d "
             "projection head, confirmed by an SVM), a 4,134-encoding enrollment store "
             "spanning 113 students, and a reproducible training pipeline for the "
             "projection head.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=8)

    add_para(doc,
             "On fourteen unseen classroom photographs the system recognizes 86% of all "
             "faces (100% on nine of those photographs) with processing latency under "
             "five seconds and no false positives. The remaining errors cluster on "
             "small, pose-extreme faces in dense classroom groups \u2014 failure modes "
             "that the teacher-review UI handles gracefully. The system is production-"
             "ready in every sense that matters for an undergraduate deployment: it is "
             "documented, tested, zero-cost to operate, privacy-respecting, and "
             "maintainable by a single developer.",
             size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.3, after=10)

    add_heading_styled(doc, "7.2 Future Recommendations", level=2)

    add_para(doc, "Several natural extensions remain:", after=4)
    add_numbered(doc, [
        "Dedicated classroom cameras. The current workflow relies on a teacher "
        "photograph. A lightweight IoT deployment \u2014 a fixed Raspberry Pi or "
        "ESP32-CAM with a 5-megapixel sensor per classroom, triggered automatically "
        "at the start of each period \u2014 would remove the last manual step from the "
        "attendance loop.",

        "Mobile-native application. A React Native or Flutter client would be "
        "particularly useful for teachers teaching in the field, in laboratories, or "
        "in overflow classrooms where browser-based upload is awkward.",

        "GPU-backed recognition. Running the same pipeline on a modest GPU (e.g. an "
        "NVIDIA T4) would reduce per-photo latency below one second and enable real-"
        "time video-stream attendance in addition to still-image attendance.",

        "Larger and more representative dataset. Expanding beyond 114 students to a "
        "faculty-wide enrollment would stress-test the projection head\u2019s ability "
        "to maintain intra-/inter-class separation at N \u2265 500.",

        "Liveness detection. Guarding against photograph-of-a-photograph attacks via "
        "a lightweight liveness check (motion, head pose, or a challenge-response) "
        "would close the one remaining security loophole in the current design.",

        "Longitudinal analytics. Because the backend already persists every "
        "attendance record, a small analytics module could surface trends at the "
        "student, course, and faculty level \u2014 useful for intervention on "
        "chronically absent students.",

        "Federated / multi-institute deployment. Nothing in the architecture prevents "
        "deploying the same stack at multiple institutes with per-tenant isolation; "
        "the model could be fine-tuned per tenant without cross-institute data "
        "sharing.",
    ])

    add_page_break(doc)


# =============================================================================
# References
# =============================================================================

def build_references(doc):
    add_heading_styled(doc, "References", level=1)
    refs = [
        "[1] F. Schroff, D. Kalenichenko, and J. Philbin, \u201CFaceNet: A Unified "
        "Embedding for Face Recognition and Clustering,\u201D Proc. IEEE Conf. Computer "
        "Vision and Pattern Recognition (CVPR), 2015.",

        "[2] K. Zhang, Z. Zhang, Z. Li, and Y. Qiao, \u201CJoint Face Detection and "
        "Alignment Using Multi-task Cascaded Convolutional Networks,\u201D IEEE Signal "
        "Processing Letters, vol. 23, no. 10, pp. 1499\u20131503, 2016.",

        "[3] J. Deng, J. Guo, N. Xue, and S. Zafeiriou, \u201CArcFace: Additive Angular "
        "Margin Loss for Deep Face Recognition,\u201D Proc. IEEE/CVF Conf. Computer "
        "Vision and Pattern Recognition (CVPR), 2019.",

        "[4] J. Deng, J. Guo, Y. Zhou, J. Yu, I. Kotsia, and S. Zafeiriou, \u201CRetinaFace: "
        "Single-stage Dense Face Localisation in the Wild,\u201D arXiv:1905.00641, 2019.",

        "[5] P. Viola and M. Jones, \u201CRapid Object Detection using a Boosted Cascade "
        "of Simple Features,\u201D Proc. IEEE Conf. Computer Vision and Pattern "
        "Recognition (CVPR), 2001.",

        "[6] T. Chen, S. Kornblith, M. Norouzi, and G. Hinton, \u201CA Simple Framework "
        "for Contrastive Learning of Visual Representations,\u201D Proc. Int. Conf. "
        "Machine Learning (ICML), 2020.",

        "[7] O. M. Parkhi, A. Vedaldi, and A. Zisserman, \u201CDeep Face Recognition,\u201D "
        "Proc. British Machine Vision Conf. (BMVC), 2015.",

        "[8] Q. Cao, L. Shen, W. Xie, O. M. Parkhi, and A. Zisserman, \u201CVGGFace2: A "
        "Dataset for Recognising Faces across Pose and Age,\u201D Proc. IEEE Int. Conf. "
        "Automatic Face and Gesture Recognition, 2018.",

        "[9] C. Cortes and V. Vapnik, \u201CSupport-Vector Networks,\u201D Machine "
        "Learning, vol. 20, no. 3, pp. 273\u2013297, 1995.",

        "[10] D. P. Kingma and J. Ba, \u201CAdam: A Method for Stochastic "
        "Optimization,\u201D Proc. Int. Conf. Learning Representations (ICLR), 2015.",

        "[11] S. Ioffe and C. Szegedy, \u201CBatch Normalization: Accelerating Deep "
        "Network Training by Reducing Internal Covariate Shift,\u201D Proc. Int. Conf. "
        "Machine Learning (ICML), 2015.",

        "[12] A. Krizhevsky, I. Sutskever, and G. E. Hinton, \u201CImageNet "
        "Classification with Deep Convolutional Neural Networks,\u201D Advances in "
        "Neural Information Processing Systems (NIPS), 2012.",

        "[13] C. Szegedy, W. Liu, Y. Jia, P. Sermanet, et al., \u201CGoing Deeper with "
        "Convolutions,\u201D Proc. IEEE Conf. Computer Vision and Pattern Recognition "
        "(CVPR), 2015.",

        "[14] K. Simonyan and A. Zisserman, \u201CVery Deep Convolutional Networks for "
        "Large-Scale Image Recognition,\u201D Proc. Int. Conf. Learning Representations "
        "(ICLR), 2015.",

        "[15] G. B. Huang, M. Mattar, T. Berg, and E. Learned-Miller, \u201CLabeled Faces "
        "in the Wild: A Database for Studying Face Recognition in Unconstrained "
        "Environments,\u201D Tech. Rep., University of Massachusetts, Amherst, 2007.",

        "[16] B. Amos, B. Ludwiczuk, and M. Satyanarayanan, \u201COpenFace: A General-"
        "Purpose Face Recognition Library with Mobile Applications,\u201D Tech. Rep., "
        "Carnegie Mellon University, 2016.",

        "[17] N. Lamba et al., \u201CSmart Attendance System using Face Recognition,\u201D "
        "Int. J. Engineering Research & Technology, vol. 10, issue 5, 2021.",

        "[18] S. Aravind et al., \u201CAttendance Monitoring System Using Face "
        "Recognition,\u201D Int. Journal of Engineering and Advanced Technology, 2019.",

        "[19] International Organization for Standardization, \u201CISO 9001:2015 "
        "Quality Management Systems \u2014 Requirements,\u201D 5th ed., 2015.",

        "[20] International Organization for Standardization, \u201CISO/IEC 27001:2022 "
        "Information Security, Cybersecurity and Privacy Protection \u2014 "
        "Information Security Management Systems \u2014 Requirements,\u201D 3rd ed., 2022.",

        "[21] United Nations, \u201CThe 17 Sustainable Development Goals,\u201D "
        "available at https://sdgs.un.org/goals, accessed April 2026.",

        "[22] Prisma Documentation, https://www.prisma.io/docs, accessed April 2026.",

        "[23] FastAPI Documentation, https://fastapi.tiangolo.com, accessed April 2026.",

        "[24] React Documentation, https://react.dev, accessed April 2026.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        run = p.add_run(r)
        run.font.size = Pt(10.5)
        run.font.name = "Times New Roman"
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_spacing(p, before=0, after=4, line=1.3)

    add_page_break(doc)


# =============================================================================
# Appendix A: Database Schema
# =============================================================================

def build_appendix_a(doc):
    add_heading_styled(doc, "Appendix A: Database Schema", level=1)
    add_para(doc,
             "The complete Prisma schema follows. Relational links are enforced at the "
             "database level; soft deletes via a `deletedAt` column are used throughout.",
             size=11, after=8, line=1.3)

    add_heading_styled(doc, "A.1 User and Authentication Tables", level=2)
    add_para(doc,
             "The User table stores both admin and teacher accounts; the role column "
             "holds either 'ADMIN' or 'TEACHER'. Passwords are bcrypt-hashed before "
             "storage. Refresh tokens are stored as hashes in the RefreshToken table "
             "with an expiration timestamp and a revocation flag.",
             size=11, after=8, line=1.3)

    add_heading_styled(doc, "A.2 Core Academic Tables", level=2)
    add_para(doc,
             "Department and Batch are small reference tables. Teacher and Student both "
             "carry a department_id foreign key; Student additionally carries a batch_id "
             "and a registration_number that doubles as a human-readable primary key.",
             size=11, after=8, line=1.3)

    add_heading_styled(doc, "A.3 Scheduling and Enrollment", level=2)
    add_para(doc,
             "The Course table holds courses offered in a given semester; "
             "TeacherCourseAssignment maps teachers to courses; StudentEnrollment maps "
             "students to courses. Classrooms are tracked in the Classroom table and "
             "per-week schedules in TimetableSlot.",
             size=11, after=8, line=1.3)

    add_heading_styled(doc, "A.4 Attendance Tables", level=2)
    add_para(doc,
             "An AttendanceSession is created per course, per date, per time slot. "
             "AttendanceRecord rows link students to sessions with a status "
             "(PRESENT/ABSENT/LATE/EXCUSED), a confidence score (if AI-recognized), and "
             "an optional override reason.",
             size=11, after=8, line=1.3)

    add_heading_styled(doc, "A.5 Audit and Derived Storage", level=2)
    add_para(doc,
             "AuditLog is append-only and records every state-changing operation "
             "(user, action, entity type, entity id, previous state, new state, "
             "timestamp, IP address). The FaceEncoding table stores only metadata "
             "pointers; the NumPy arrays themselves live on disk under "
             "`ai-service/encodings/<student-uuid>/`.",
             size=11, after=8, line=1.3)

    add_page_break(doc)


# =============================================================================
# Appendix B: API Reference (condensed)
# =============================================================================

def build_appendix_b(doc):
    add_heading_styled(doc, "Appendix B: API Reference (Condensed)", level=1)
    add_para(doc,
             "The full REST reference is available in the repository under "
             "`backend/docs/api.md`; a condensed summary follows.",
             size=11, after=8, line=1.3)

    add_heading_styled(doc, "B.1 Authentication", level=2)
    add_bullets(doc, [
        "POST /api/v1/auth/login \u2014 body: { email, password }. Returns "
        "{ accessToken, refreshToken, user }.",
        "POST /api/v1/auth/refresh \u2014 body: { refreshToken }. Returns a new "
        "access token.",
        "POST /api/v1/auth/logout \u2014 invalidates the caller\u2019s refresh token.",
        "GET /api/v1/auth/profile \u2014 returns current user.",
        "POST /api/v1/auth/change-password \u2014 body: { oldPassword, newPassword }.",
    ])

    add_heading_styled(doc, "B.2 Admin Endpoints", level=2)
    add_bullets(doc, [
        "Students: /api/v1/admin/students \u2014 standard CRUD.",
        "Teachers: /api/v1/admin/teachers \u2014 standard CRUD.",
        "Courses: /api/v1/admin/courses \u2014 standard CRUD.",
        "Classrooms: /api/v1/admin/classrooms \u2014 standard CRUD.",
        "Timetables: /api/v1/admin/timetables \u2014 standard CRUD.",
        "Audit logs: GET /api/v1/admin/logs/audit, GET /api/v1/admin/logs/attendance.",
    ])

    add_heading_styled(doc, "B.3 Teacher Endpoints", level=2)
    add_bullets(doc, [
        "GET /api/v1/teacher/schedule \u2014 returns the teacher\u2019s weekly schedule.",
        "POST /api/v1/teacher/sessions \u2014 creates a session for a course/date/slot.",
        "POST /api/v1/teacher/sessions/:id/image \u2014 multipart upload of the "
        "classroom photograph; returns recognized students + unknown faces.",
        "GET /api/v1/teacher/sessions/:id \u2014 returns session state with records.",
        "POST /api/v1/teacher/sessions/:id/finalize \u2014 commits the attendance "
        "records; after finalization the session is read-only.",
    ])

    add_heading_styled(doc, "B.4 AI Service Endpoints", level=2)
    add_bullets(doc, [
        "POST /api/v1/recognize \u2014 body: multipart { image, student_ids, threshold }. "
        "Returns { facesDetected, facesRecognized, recognizedStudents[], unknownFaces[], "
        "metrics{}, annotatedImageBase64 }.",
        "POST /api/v1/encodings \u2014 body: multipart { student_id, images[] }. "
        "Enrolls or updates a student\u2019s encoding store.",
        "DELETE /api/v1/encodings/:studentId \u2014 removes a student\u2019s encodings.",
        "GET /api/v1/health \u2014 readiness probe; returns model load status.",
    ])

    add_page_break(doc)


# =============================================================================
# Appendix C: Training Configuration
# =============================================================================

def build_appendix_c(doc):
    add_heading_styled(doc, "Appendix C: Training Configuration", level=1)

    add_heading_styled(doc, "C.1 Projection Head Training", level=2)
    add_para(doc,
             "The projection head is trained by `train_projection_head.py`. Default "
             "hyperparameters:",
             size=11, after=6, line=1.3)

    add_table(doc,
              headers=["Parameter", "Value"],
              rows=[
                  ["Architecture", "512 \u2192 256 \u2192 128 dense, ReLU, BN, dropout 0.1"],
                  ["Final layer", "L2-normalisation"],
                  ["Loss function", "Triplet loss, margin \u03b1 = 0.3"],
                  ["Optimizer", "Adam, learning rate 1e-3"],
                  ["Batch size (triplets)", "64"],
                  ["Steps per epoch", "50"],
                  ["Max epochs", "80"],
                  ["Early stopping", "Patience 15 (on val triplet loss)"],
                  ["Train/val split", "By student ID, 80/20"],
                  ["Noise augmentation", "\u03c3 = 0.01 on anchor/positive, re-normalised"],
              ],
              caption="",
              col_widths=[2.5, 4.0])

    add_heading_styled(doc, "C.2 SVM Training", level=2)
    add_para(doc,
             "The SVM is trained by `retrain_svm.py`. Hyperparameters are chosen by "
             "three-fold grid search when at least 50 samples and 5 classes are "
             "available; otherwise fixed defaults are used.",
             size=11, after=6, line=1.3)

    add_table(doc,
              headers=["Parameter", "Default / Grid"],
              rows=[
                  ["Kernel", "RBF"],
                  ["Probability", "True"],
                  ["Class weight", "Balanced"],
                  ["Grid search C", "{1, 10, 50, 100}"],
                  ["Grid search gamma", "{scale, auto}"],
                  ["CV folds", "min(3, n_classes)"],
                  ["Fixed fallback", "C = 10.0, gamma = scale"],
              ],
              caption="",
              col_widths=[2.5, 4.0])

    add_heading_styled(doc, "C.3 Enrichment Script", level=2)
    add_para(doc,
             "`enrich_from_groups.py` walks a folder of classroom group photographs and "
             "adds high-confidence classroom face crops to each student\u2019s encoding "
             "store. Guards applied:",
             size=11, after=6, line=1.3)
    add_bullets(doc, [
        "Tight L2 threshold for the initial match: 0.75 in raw 512-d FaceNet space.",
        "Minimum margin to second-best match: 0.10.",
        "Pass through the same quality filter used during recognition.",
        "After enrichment, recompute centroids and stats, re-project, and retrain SVM.",
    ])

    add_page_break(doc)


# =============================================================================
# Appendix D: Deployment Notes
# =============================================================================

def build_appendix_d(doc):
    add_heading_styled(doc, "Appendix D: Deployment Notes", level=1)

    add_para(doc,
             "The target deployment is a single Linux VM hosting all three services "
             "behind an nginx reverse proxy that terminates TLS. Oracle Cloud\u2019s "
             "always-free Ampere A1 instance (4 OCPU, 24 GB RAM, 200 GB disk) is the "
             "reference host because it satisfies the AI service\u2019s memory "
             "requirements at zero recurring cost.",
             size=11, after=8, line=1.3, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading_styled(doc, "D.1 Topology", level=2)
    add_para(doc,
             "nginx listens on ports 80 and 443 and routes requests by path: `/` to the "
             "React static bundle, `/api/v1/` to the Node.js backend on port 3000, and "
             "internal backend-to-AI calls to the FastAPI service on localhost port 8000. "
             "TLS certificates are provisioned via certbot + Let\u2019s Encrypt against a "
             "DuckDNS subdomain.",
             size=11, after=8, line=1.3, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading_styled(doc, "D.2 Process Supervision", level=2)
    add_para(doc,
             "Both the backend and the AI service run under systemd unit files that "
             "specify `Restart=on-failure` with a five-second back-off. Logs go to the "
             "journal and are rotated daily; access logs from nginx are shipped to the "
             "same journal for unified retention.",
             size=11, after=8, line=1.3, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading_styled(doc, "D.3 Backups", level=2)
    add_para(doc,
             "The SQLite database is snapshotted nightly to an encrypted tarball via a "
             "systemd timer; seven daily rolls are retained, plus one monthly archive. "
             "Encoding files are snapshotted weekly (they change rarely). The `.env` "
             "file containing JWT secrets is not included in backups and must be "
             "regenerated on restore.",
             size=11, after=8, line=1.3, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_heading_styled(doc, "D.4 Restore Procedure", level=2)
    add_numbered(doc, [
        "Provision a fresh VM with Ubuntu 22.04.",
        "Install Node.js 18, Python 3.10+, and system libraries (libgl1, libglib2.0).",
        "Clone the repository, install npm and pip dependencies.",
        "Restore the most recent database tarball; run Prisma migrations.",
        "Restore the encoding directory; verify `get_all_student_ids()` reports the "
        "expected count.",
        "Generate fresh JWT secrets; seed admin account via `prisma db seed`.",
        "Start systemd services; verify /api/v1/health returns models-loaded.",
    ])

    add_page_break(doc)


# =============================================================================
# Appendix E: User Guide
# =============================================================================

def build_appendix_e(doc):
    add_heading_styled(doc, "Appendix E: User Guide (Quick Reference)", level=1)

    add_heading_styled(doc, "E.1 Admin Workflow", level=2)
    add_numbered(doc, [
        "Log in with the admin credentials.",
        "Add departments and batches.",
        "Upload students (CSV import or row-by-row entry).",
        "Create teacher accounts and assign them to courses.",
        "Define classrooms and the weekly timetable.",
        "Invite students to submit 3\u20135 selfie photographs for enrollment.",
        "Run `generate_encodings.py` to populate the encoding store.",
        "Run `train_projection_head.py` and `retrain_svm.py`.",
        "Monitor `/admin/logs/audit` and `/admin/logs/attendance` as required.",
    ])

    add_heading_styled(doc, "E.2 Teacher Workflow", level=2)
    add_numbered(doc, [
        "Log in with teacher credentials.",
        "Open the Schedule view; click the slot for the current period.",
        "On the Session page, take or upload a classroom photograph.",
        "Wait 1\u20135 seconds; inspect the annotated result.",
        "Correct any misrecognitions and click Finalize.",
        "Review past sessions under Attendance History at any time.",
    ])

    add_heading_styled(doc, "E.3 Troubleshooting", level=2)
    add_bullets(doc, [
        "No faces detected: ensure the photograph is in focus and faces are at "
        "least 20 px in the shortest dimension.",
        "Many unknowns: a student may not be enrolled, or their encodings may be "
        "stale \u2014 re-enroll them.",
        "Slow processing: check the image resolution; very large images bypass the "
        "resize only if MAX_DIM has been customised.",
        "Recognition service unreachable: the backend will store the image and "
        "queue the session for reprocessing automatically.",
    ])

    add_page_break(doc)


# =============================================================================
# Main
# =============================================================================

def main():
    doc = setup_document()

    build_cover(doc)
    build_dedication(doc)
    build_certificate(doc)
    build_preface(doc)
    build_acknowledgments(doc)
    build_toc(doc)
    build_lof(doc)
    build_lot(doc)
    build_abstract(doc)
    build_cep(doc)
    build_chapter_1(doc)
    build_chapter_2(doc)
    build_chapter_3(doc)
    build_chapter_4(doc)
    build_chapter_5(doc)
    build_chapter_6(doc)
    build_chapter_7(doc)
    build_references(doc)
    build_appendix_a(doc)
    build_appendix_b(doc)
    build_appendix_c(doc)
    build_appendix_d(doc)
    build_appendix_e(doc)

    doc.save(OUT)
    print(f"Report saved: {OUT}")


if __name__ == "__main__":
    main()
